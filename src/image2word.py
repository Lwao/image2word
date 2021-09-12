# -*- coding: utf-8 -*-
"""
Created on Sun Jun 21 18:20:22 2020

@author: levyg
"""

# pyinstaller .\faturas.py
# pyinstaller -c -F -i cm_icon.ico ./image2word.py
# pyinstaller -c -F -i ./faul.ico ./image2word.py
# pyinstaller -c -F -i ./sloth.ico ./image2word.py


from docx import Document
from docx.shared import Inches
from tkinter import filedialog
from docxcompose.composer import Composer
import tkinter as tk
import os
import os.path


print("Organizador de imagens\nver.: 1.0.1, 30/09/2020\nAutor: Levy Gabriel")
input("Pressione ENTER para iniciar...")
print("\nSelecione a pasta que contém as pastas dos alunos:")

try:
    # load files paths
    root = tk.Tk()
    root.withdraw()
    root.call('wm', 'attributes', '.', '-topmost', True)
    files = filedialog.askdirectory()
    directory = files
    directory_name = os.path.basename(os.path.normpath(directory))
    print(directory_name)
    subdirectory = os.listdir(directory) 
    images_per_dir = []
    tot_images = []
    for student_directory in subdirectory:
        images_per_dir.append(os.listdir(directory + "/" + student_directory))
        tot_images.append(len(images_per_dir[len(images_per_dir)-1][:]))
        
    print("\nSelecione o documento do plano de aula:")
    
    # first document that is already generated (plano de aula)
    merged_word_address = filedialog.askopenfile()
    directory_word = merged_word_address.name
    print(os.path.basename(os.path.normpath(directory_word)))
    #append_word_document = Document(append_word_address.name)
    #append_word_document.add_page_break()
    
    # second document with images
    document = Document(merged_word_address.name)    
    document.add_page_break()
    #document.add_heading('Plano de aula:')
    
    #p = document.add_paragraph('Constam em anexo imagens referente ')
    #p.add_run('bold').bold = True
    #p.add_run(' and some ')
    #p.add_run('italic.').italic = True
    
    for i in range(len(subdirectory)):
        qnt_images = tot_images[i]
        
        document.add_paragraph(str(i+1) + '. ' + subdirectory[i])
    #    document.add_paragraph(subdirectory[i], style='List Number')
        tbl = document.add_table(rows=1, cols=2)
        for j in range(qnt_images):
            if (((j+1)%2)!=0):
                row_cells = tbl.add_row().cells
                paragraph = row_cells[0].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(directory + "/" + subdirectory[i] + "/" + images_per_dir[i][j] , width = Inches(3), height = Inches(3))
            elif (((j+1)%2)==0):
                paragraph = row_cells[1].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(directory + "/" + subdirectory[i] + "/" + images_per_dir[i][j], width = Inches(3), height = Inches(3))
        document.add_page_break()
    
    # append documents
    #merged = Composer(append_word_document)
    #merged.append(document)
    #
    
    #composer.save(directory[:len(directory)-len(directory_name)] + 'resultado.docx')
    
    document.save(directory[:len(directory)-len(directory_name)] + 'resultado.docx')


    print("\nO arquivo resultado.docx foi salvo com sucesso ao lado da pasta: " + directory_name + ".")
    input("Pressione ENTER para finalizar...")
except:
    print("\nOcorreu um erro ao selecionar as pastas. Por favor, tente novamente com o procedimento correto!")
    input("Pressione ENTER para finalizar...")